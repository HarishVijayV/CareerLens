"""
Turn an uploaded resume file into plain text (and LaTeX, when that's what was uploaded).

Format-by-format reality:

  .tex   — best case. It's already text, AND it's the thing that compiles to the PDF you
           actually send, so an AI edit can flow all the way back to a real document.
           We keep the raw LaTeX *and* a stripped-down text view for the agents.

  .pdf   — extractable, but LOSSY and one-way. PDF stores glyphs and positions, not
           structure; two-column layouts interleave, and there is no reliable path back
           to a formatted PDF. Fine for reading and matching, not for round-tripping.

  .docx  — a zip of XML. python-docx pulls paragraphs and tables out cleanly enough.

  .txt/.md — nothing to do.

The honest hierarchy: upload .tex if you have it, because it's the only format where an
AI edit can produce a document you can actually send.
"""
import io
import re
import zipfile


class ResumeParseError(Exception):
    pass


def extract_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ResumeParseError("pypdf not installed — cannot read PDF uploads")

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ResumeParseError(f"Could not read PDF: {exc}")

    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()

    if not text:
        # A scanned/image-only PDF has no text layer at all. Say so plainly — this looks
        # identical to "upload failed" from the outside, and the fix (OCR, or paste the
        # text) is completely different.
        raise ResumeParseError(
            "No text found in this PDF. It may be a scan or image export — "
            "paste the text directly, or upload the .tex/.docx source instead."
        )
    return text


def extract_from_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise ResumeParseError("python-docx not installed — cannot read .docx uploads")

    try:
        document = docx.Document(io.BytesIO(data))
    except (zipfile.BadZipFile, Exception) as exc:
        raise ResumeParseError(f"Could not read .docx: {exc}")

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Resumes very often lay out dates/sections in tables, and those paragraphs are NOT
    # in document.paragraphs — miss this and half the resume silently disappears.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))

    return "\n".join(parts).strip()


def latex_to_text(latex: str) -> str:
    """Strip LaTeX markup down to readable text for the agents.

    Deliberately a pragmatic regex pass, not a real parser: the goal is legible content
    for an LLM to reason about, not a faithful re-render. Feeding raw LaTeX to the model
    instead would burn tokens on syntax and tempt it to "fix" formatting it shouldn't
    touch.
    """
    text = latex

    text = re.sub(r"(?<!\\)%.*", "", text)                       # comments (not \%)
    text = re.sub(r"\\(usepackage|documentclass|input|include|hypersetup|geometry)"
                  r"(\[[^\]]*\])?\{[^}]*\}", "", text)           # preamble noise
    text = re.sub(r"\\(begin|end)\{[^}]*\}", "\n", text)          # environments
    text = re.sub(r"\\(section|subsection|subsubsection)\*?\{([^}]*)\}", r"\n\n\2\n", text)
    text = re.sub(r"\\(textbf|textit|emph|underline|texttt|large|Large|href)"
                  r"(\{[^}]*\})?\{([^}]*)\}", r"\3", text)        # keep the visible arg
    text = re.sub(r"\\item\s*", "- ", text)
    text = re.sub(r"\\\\|\\newline|\\par", "\n", text)
    # Inline math is usually a separator glyph in a resume ($\cdot$, $\bullet$). Replace
    # the whole span with a bullet rather than stripping only the command, which would
    # otherwise leave stray "$$" in the text handed to the agents.
    text = re.sub(r"\$[^$]*\$", "·", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{[^}]*\})?", "", text)  # leftover commands
    text = text.replace("~", " ").replace("\\&", "&").replace("\\%", "%").replace("\\_", "_")
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)

    return text.strip()


def parse_upload(filename: str, data: bytes) -> tuple[str, str | None, str]:
    """Returns (plain_text, latex_or_none, detected_format)."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        return extract_from_pdf(data), None, "pdf"

    if name.endswith((".tex", ".latex")):
        latex = data.decode("utf-8", errors="replace")
        return latex_to_text(latex), latex, "tex"

    if name.endswith(".docx"):
        return extract_from_docx(data), None, "docx"

    if name.endswith((".txt", ".md", ".markdown")):
        return data.decode("utf-8", errors="replace").strip(), None, "txt"

    if name.endswith(".doc"):
        # Legacy binary .doc is a different format entirely from .docx and python-docx
        # cannot read it. Better to say so than to return garbled bytes.
        raise ResumeParseError(
            "Legacy .doc isn't supported. Save as .docx or .pdf and upload again."
        )

    raise ResumeParseError(f"Unsupported file type: {filename}. Use .tex, .pdf, .docx, or .txt")
